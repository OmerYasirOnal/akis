# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN, WD_TAB_ALIGNMENT as TABAL, WD_TAB_LEADER as LEAD
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC="/tmp/report/build/working.docx"
OUT="/tmp/report/build/final.docx"
CH="/tmp/report/build/charts"
doc=Document(SRC)
body=doc.element.body
FONT="Times New Roman"

def mk(tag): return OxmlElement('w:'+tag)
def setval(el,attr,val): el.set(qn('w:'+attr),val); return el

def make_run(text, bold=False, italic=False, sz=24, color=None):
    r=mk('r'); rpr=mk('rPr')
    rf=mk('rFonts'); rf.set(qn('w:ascii'),FONT); rf.set(qn('w:hAnsi'),FONT); rf.set(qn('w:cs'),FONT); rpr.append(rf)
    if bold: rpr.append(mk('b')); rpr.append(mk('bCs'))
    if italic: rpr.append(mk('i')); rpr.append(mk('iCs'))
    sz_el=mk('sz'); setval(sz_el,'val',str(sz)); rpr.append(sz_el)
    szcs=mk('szCs'); setval(szcs,'val',str(sz)); rpr.append(szcs)
    if color: c=mk('color'); setval(c,'val',color); rpr.append(c)
    r.append(rpr)
    t=mk('t'); t.set(qn('xml:space'),'preserve'); t.text=text; r.append(t)
    return r

_BM=[9000]
def bookmark_start(name):
    b=mk('bookmarkStart'); b.set(qn('w:id'),str(_BM[0])); b.set(qn('w:name'),name); return b
def bookmark_end():
    b=mk('bookmarkEnd'); b.set(qn('w:id'),str(_BM[0])); _BM[0]+=1; return b

def field_runs(instr, cached, bold=False, italic=False, sz=24):
    out=[]
    r=mk('r'); fc=mk('fldChar'); setval(fc,'fldCharType','begin'); r.append(fc); out.append(r)
    r=mk('r'); it=mk('instrText'); it.set(qn('xml:space'),'preserve'); it.text=instr; r.append(it); out.append(r)
    r=mk('r'); fc=mk('fldChar'); setval(fc,'fldCharType','separate'); r.append(fc); out.append(r)
    out.append(make_run(cached,bold=bold,italic=italic,sz=sz))
    r=mk('r'); fc=mk('fldChar'); setval(fc,'fldCharType','end'); r.append(fc); out.append(r)
    return out

def new_p(style_jc='center', spacing_before=None, spacing_after=None, line=None):
    p=mk('p'); ppr=mk('pPr')
    if spacing_before is not None or spacing_after is not None or line is not None:
        sp=mk('spacing')
        if spacing_before is not None: setval(sp,'before',str(spacing_before))
        if spacing_after is not None: setval(sp,'after',str(spacing_after))
        if line is not None: setval(sp,'line',str(line)); setval(sp,'lineRule','auto')
        ppr.append(sp)
    jc=mk('jc'); setval(jc,'val',style_jc); ppr.append(jc)
    p.append(ppr); return p

def caption_p(label, rest, bm, italic_rest):
    # label e.g. "Figure 6.1" (bold, bookmarked); rest e.g. "AKIS multi-agent ..." 
    p=new_p('center', spacing_before=120, spacing_after=240, line=240)  # fig: 6pt before,12pt after
    if not italic_rest:  # table caption: 12 before, 6 after
        p.find(qn('w:pPr')+'/'+qn('w:spacing')).set(qn('w:before'),'240')
        p.find(qn('w:pPr')+'/'+qn('w:spacing')).set(qn('w:after'),'120')
    p.append(bookmark_start(bm))
    p.append(make_run(label, bold=True, sz=22))
    p.append(bookmark_end())
    p.append(make_run(". ", bold=True, sz=22))
    p.append(make_run(rest, italic=italic_rest, sz=22))
    return p

def add_pic_elem(img_path, width_in):
    p=doc.add_paragraph(); p.alignment=ALIGN.CENTER
    p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(0)
    p.add_run().add_picture(img_path, width=Inches(width_in))
    return p._p  # currently appended at end of body; will be relocated

def text_of(p):
    return ''.join(t.text for t in p.findall('.//'+qn('w:t')) if t.text)

def style_of(p):
    s=p.find(qn('w:pPr')+'/'+qn('w:pStyle'))
    return s.get(qn('w:val')) if s is not None else "normal"

paras=[blk for blk in body.iterchildren() if blk.tag==qn('w:p')]
def find_p(pred):
    for p in body.iterchildren():
        if p.tag==qn('w:p') and pred(p): return p
    return None
def find_by_text(sub, style=None, exact_start=False):
    for p in body.iterchildren():
        if p.tag!=qn('w:p'): continue
        tx=text_of(p).strip()
        ok = tx.startswith(sub) if exact_start else (sub in tx)
        if ok and (style is None or style_of(p)==style):
            return p
    return None

def delete_p(p):
    if p is not None: p.getparent().remove(p)

def clear_runs(p):
    for r in list(p):
        if r.tag in (qn('w:r'),qn('w:bookmarkStart'),qn('w:bookmarkEnd'),qn('w:hyperlink')):
            p.remove(r)

def rebuild_caption(p, label, rest, bm, italic_rest, keep_image=False):
    # keep pPr; optionally keep first drawing run; replace text content with bold label + rest, add bookmark
    img_run=None
    if keep_image:
        for r in p.findall(qn('w:r')):
            if r.find('.//'+qn('w:drawing')) is not None: img_run=r; break
    clear_runs(p)
    # rebuild pPr cleanly in schema order: spacing, jc  (drop old pStyle/rPr)
    old=p.find(qn('w:pPr'))
    if old is not None: p.remove(old)
    ppr=mk('pPr')
    sp=mk('spacing')
    if italic_rest: setval(sp,'before','120'); setval(sp,'after','240')
    else: setval(sp,'before','240'); setval(sp,'after','120')
    setval(sp,'line','240'); setval(sp,'lineRule','auto'); ppr.append(sp)
    jc=mk('jc'); setval(jc,'val','center'); ppr.append(jc)
    p.insert(0,ppr)
    if keep_image and img_run is not None:
        p.append(img_run)
        p.append(mk('r')); p[-1].append(mk('br'))  # line break after image
    p.append(bookmark_start(bm))
    p.append(make_run(label, bold=True, sz=22))
    p.append(bookmark_end())
    p.append(make_run(". ", bold=True, sz=22))
    p.append(make_run(rest, italic=italic_rest, sz=22))
    return p

print("== STEP 1: split merged caption+heading paragraphs ==")
# [352] Figure 6.1 caption merged with heading "6.2 AKIS Workflow Orchestrator"
p61=find_by_text("Figure 6.1. AKIS multi-agent")
# Extract heading text from same paragraph, create separate heading paragraph AFTER caption
def split_off_heading(p_caption, heading_text):
    new_h=mk('p'); ppr=mk('pPr'); pst=mk('pStyle'); setval(pst,'val','Heading2'); ppr.append(pst); new_h.append(ppr)
    new_h.append(make_run(heading_text, bold=True, sz=26))
    p_caption.addnext(new_h)
    return new_h
split_off_heading(p61, "6.2 AKIS Workflow Orchestrator")
rebuild_caption(p61,"Figure 6.1",
    "AKIS multi-agent workflow pipeline. Critic is invoked at two distinct stages — after specification generation and after code generation. Trace retrieves committed code directly from GitHub, ensuring independent verification.",
    "_Ref_fig_6_1", italic_rest=True)

print("== STEP 2: figure caption renumber + dedup ==")
# 6.2 human approval [401]
delete_p(find_by_text("Figure 6.X. Human approval gate before"))  # short dup if exists
p62=find_by_text("Figure 6.X. Human approval checkpoint")
rebuild_caption(p62,"Figure 6.2",
    "Human approval checkpoint ensuring user oversight before transitioning from specification generation to implementation activities.",
    "_Ref_fig_6_2", italic_rest=True)
# 6.3 trust ledger: delete the before-image caption [412], keep after-image [414]
delete_p(find_by_text("Figure 6.X. Trust Ledger recording"))
p63=find_by_text("Figure 6.X. Trust Ledger displaying")
rebuild_caption(p63,"Figure 6.3",
    "Trust Ledger displaying workflow stages, approval checkpoints, and execution status throughout the software generation process.",
    "_Ref_fig_6_3", italic_rest=True)
# 7.1 delete heading3 short, keep long
delete_p(find_by_text("Figure 7.1. AKIS Studio Interface", style="Heading3"))
p71=find_by_text("Figure 7.1. AKIS Studio interface where")
rebuild_caption(p71,"Figure 7.1",
    "AKIS Studio interface where users describe software requirements and initiate agent workflows.",
    "_Ref_fig_7_1", italic_rest=True)
# 7.2
delete_p(find_by_text("Figure 7.2.Multi-Agent Workflow Execution", style="Heading3"))
p72=find_by_text("Figure 7.2. Execution of the multi-agent")
rebuild_caption(p72,"Figure 7.2",
    "Execution of the multi-agent workflow showing implementation, verification, and approval stages.",
    "_Ref_fig_7_2", italic_rest=True)
# 7.3 clarification: delete short [495]; [496] = img+caption+heading "7.6 Verification Features"
delete_p(find_by_text("Figure 7.3. Requirement clarification questions generated"))
p73=find_by_text("Figure 7.3. Requirement clarification phase")
# split heading 7.6 out of p73 (keep image)
split_off_heading(p73, "7.6 Verification Features")
rebuild_caption(p73,"Figure 7.3",
    "Requirement clarification phase where the Scribe agent generates follow-up questions to resolve ambiguities and improve specification quality before user approval.",
    "_Ref_fig_7_3", italic_rest=True, keep_image=True)
# 7.4 generated app: delete heading3 short, keep long, renumber 7.3->7.4
delete_p(find_by_text("Figure 7.3.Generated Application Example", style="Heading3"))
p74=find_by_text("Figure 7.3. Example note-taking application")
rebuild_caption(p74,"Figure 7.4",
    "Example note-taking application generated through the AKIS workflow.",
    "_Ref_fig_7_4", italic_rest=True)
# 8.1, 8.2
p81=find_by_text("Figure 8.1. Trust Report interface")
rebuild_caption(p81,"Figure 8.1",
    "Trust Report interface presenting verification evidence, executed test scenarios, approval status, and validation outcomes. The QR code generator example run shows 2 tests passed, Critic approved with 6 findings, and deploy authorised.",
    "_Ref_fig_8_1", italic_rest=True)
p82=find_by_text("Figure 8.2. Verification dashboard")
rebuild_caption(p82,"Figure 8.2",
    "Verification dashboard summarising test execution statistics, validation status, and deployment readiness. Results shown are from a mock pipeline baseline run.",
    "_Ref_fig_8_2", italic_rest=True)

print("== STEP 3: table captions rebuild (bold label, period, bookmark, de-heading) ==")
rebuild_caption(find_by_text("Table 5.1 Functional Requirements"),"Table 5.1","Functional Requirements.","_Ref_tbl_5_1",italic_rest=False)
rebuild_caption(find_by_text("Table 5.2 Non-Functional Requirements"),"Table 5.2","Non-Functional Requirements.","_Ref_tbl_5_2",italic_rest=False)
rebuild_caption(find_by_text("Table 8.1. Summary of automated"),"Table 8.1","Summary of automated test execution results.","_Ref_tbl_8_1",italic_rest=False)
rebuild_caption(find_by_text("Table 8.2. Frontend coverage results"),"Table 8.2","Frontend coverage results.","_Ref_tbl_8_2",italic_rest=False)
rebuild_caption(find_by_text("Table 8.3. Mock pipeline baseline"),"Table 8.3","Mock pipeline baseline evaluation scenarios.","_Ref_tbl_8_3",italic_rest=False)
rebuild_caption(find_by_text("Table 9.1. Comparison of project objectives"),"Table 9.1","Comparison of project objectives and achieved results.","_Ref_tbl_9_1",italic_rest=False)

def prose_p(text, after=120, line=360):
    p=mk('p'); ppr=mk('pPr')
    sp=mk('spacing'); setval(sp,'after',str(after)); setval(sp,'line',str(line)); setval(sp,'lineRule','auto'); ppr.append(sp)
    jc=mk('jc'); setval(jc,'val','both'); ppr.append(jc)
    p.append(ppr); p.append(make_run(text, sz=24)); return p

def chart_para(img_path, width):
    return add_pic_elem(img_path, width)

print("== STEP 4: insert charts (image + caption [+ prose]) ==")
# Fig 4.1 (SWE-bench) at end of 4.6, before 4.7 heading
h47=find_by_text("4.7 Trust, Validation")
pr=prose_p("Beyond test generation, recent benchmarks have measured how reliably autonomous agents resolve real software tasks. On SWE-bench, a benchmark of real GitHub issues, the best evaluated language model (Claude 2) resolved only 1.96% of issues [16], whereas agent-based approaches such as SWE-agent raised this figure to 12.47% [4]. A subsequent analysis (SWE-Bench+) showed, however, that a substantial fraction of accepted patches relied on solutions leaked in the issue text or were admitted by weak tests; after filtering these cases the resolved rate of the same agent fell to 3.97% [17]. As summarised in Figure 4.1, autonomous resolution of real-world software issues remains far below the level required for unsupervised deployment, motivating the independent verification of generated artifacts that the AKIS Trace agent provides.")
h47.addprevious(pr)
img=chart_para(f"{CH}/fig_4_1_swebench.png",5.1); pr.addnext(img)
cap=caption_p("Figure 4.1","Reported reliability of autonomous AI agents at resolving real-world software issues on SWE-bench [4], [16], [17].","_Ref_fig_4_1",italic_rest=True)
img.addnext(cap)

# Fig 4.2 (self-correction) at end of 4.7, before 4.8 heading
h48=find_by_text("4.8 Literature Gap")
pr=prose_p("A further motivation for independent verification arises from the limited ability of language models to validate their own output. Huang et al. evaluated intrinsic self-correction — where a model revises its own answers without external feedback — and found that accuracy does not improve and frequently degrades after self-correction [18]. As shown in Figure 4.2, for GPT-3.5 the accuracy on CommonSenseQA fell from 75.8% to 38.1% after a single self-correction round, while GSM8K accuracy for both GPT-3.5 and GPT-4 declined across correction rounds. This evidence supports a core design decision of AKIS: no agent validates its own output, and verification is always performed by an independent entity.")
h48.addprevious(pr)
img=chart_para(f"{CH}/fig_4_2_selfcorrect.png",5.5); pr.addnext(img)
cap=caption_p("Figure 4.2","Reasoning accuracy before and after intrinsic self-correction, showing no improvement and frequent degradation [18].","_Ref_fig_4_2",italic_rest=True)
img.addnext(cap)

# Fig 8.3 (tests) before 8.2.2 heading
h822=find_by_text("8.2.2 Frontend Coverage Results")
pr=prose_p("Figure 8.3 visualises the distribution of the executed automated tests across the backend and frontend suites.")
h822.addprevious(pr)
img=chart_para(f"{CH}/fig_8_3_tests.png",5.2); pr.addnext(img)
cap=caption_p("Figure 8.3","Distribution of executed automated tests across backend and frontend suites (1,700 tests, 100% pass rate).","_Ref_fig_8_3",italic_rest=True)
img.addnext(cap)

# Fig 8.4 (coverage) before 8.2.3
h823=find_by_text("8.2.3 Test Category Coverage")
pr=prose_p("Figure 8.4 presents the frontend coverage metrics against the 80% target threshold.")
h823.addprevious(pr)
img=chart_para(f"{CH}/fig_8_4_coverage.png",5.2); pr.addnext(img)
cap=caption_p("Figure 8.4","Frontend code coverage by metric, measured with Vitest (v8 provider), against the 80% target.","_Ref_fig_8_4",italic_rest=True)
img.addnext(cap)

# Fig 9.1 (objectives) before 9.2 heading
h92=find_by_text("9.2 Verification and Trust Evaluation")
pr=prose_p("Figure 9.1 summarises the overall achievement of the project objectives listed in Table 9.1.")
h92.addprevious(pr)
img=chart_para(f"{CH}/fig_9_1_objectives.png",3.2); pr.addnext(img)
cap=caption_p("Figure 9.1","Achievement of project objectives (10 of 10 implemented).","_Ref_fig_9_1",italic_rest=True)
img.addnext(cap)

print("== STEP 5: List of Tables + List of Figures (before ABSTRACT) ==")
def title_heading(text, page_break=True):
    p=mk('p'); ppr=mk('pPr')
    pst=mk('pStyle'); setval(pst,'val','Title'); ppr.append(pst)
    if page_break: ppr.append(mk('pageBreakBefore'))
    sp=mk('spacing'); setval(sp,'before','240'); setval(sp,'after','240'); ppr.append(sp)
    jc=mk('jc'); setval(jc,'val','center'); ppr.append(jc)
    p.append(ppr); p.append(make_run(text, bold=True, sz=28)); return p

def list_entry(label, title, bm):
    p=mk('p'); ppr=mk('pPr')
    tabs=mk('tabs'); tb=mk('tab'); setval(tb,'val','right'); setval(tb,'leader','dot'); setval(tb,'pos','9072'); tabs.append(tb); ppr.append(tabs)
    sp=mk('spacing'); setval(sp,'before','60'); setval(sp,'after','60'); setval(sp,'line','240'); setval(sp,'lineRule','auto'); ppr.append(sp)
    jc=mk('jc'); setval(jc,'val','left'); ppr.append(jc)
    p.append(ppr)
    p.append(make_run(f"{label}. {title}", sz=22))
    tr=mk('r'); tr.append(mk('tab')); p.append(tr)
    for f in field_runs(' PAGEREF %s \\h '%bm, '0', sz=22): p.append(f)
    return p

abstract=find_by_text("ABSTRACT", style="Title")
FIG_LIST=[("Figure 4.1","Reported reliability of autonomous AI agents on SWE-bench.","_Ref_fig_4_1"),
("Figure 4.2","Reasoning accuracy before and after intrinsic self-correction.","_Ref_fig_4_2"),
("Figure 6.1","AKIS multi-agent workflow pipeline.","_Ref_fig_6_1"),
("Figure 6.2","Human approval checkpoint before implementation.","_Ref_fig_6_2"),
("Figure 6.3","Trust Ledger displaying workflow stages and approvals.","_Ref_fig_6_3"),
("Figure 7.1","AKIS Studio interface.","_Ref_fig_7_1"),
("Figure 7.2","Execution of the multi-agent workflow.","_Ref_fig_7_2"),
("Figure 7.3","Requirement clarification phase by the Scribe agent.","_Ref_fig_7_3"),
("Figure 7.4","Example note-taking application generated by AKIS.","_Ref_fig_7_4"),
("Figure 8.1","Trust Report interface.","_Ref_fig_8_1"),
("Figure 8.2","Verification dashboard.","_Ref_fig_8_2"),
("Figure 8.3","Distribution of executed automated tests.","_Ref_fig_8_3"),
("Figure 8.4","Frontend code coverage by metric.","_Ref_fig_8_4"),
("Figure 9.1","Achievement of project objectives.","_Ref_fig_9_1")]
TAB_LIST=[("Table 5.1","Functional Requirements.","_Ref_tbl_5_1"),
("Table 5.2","Non-Functional Requirements.","_Ref_tbl_5_2"),
("Table 8.1","Summary of automated test execution results.","_Ref_tbl_8_1"),
("Table 8.2","Frontend coverage results.","_Ref_tbl_8_2"),
("Table 8.3","Mock pipeline baseline evaluation scenarios.","_Ref_tbl_8_3"),
("Table 9.1","Comparison of project objectives and achieved results.","_Ref_tbl_9_1")]
blocks_to_insert=[title_heading("LIST OF TABLES")]
for lbl,ti,bm in TAB_LIST: blocks_to_insert.append(list_entry(lbl,ti,bm))
blocks_to_insert.append(title_heading("LIST OF FIGURES"))
for lbl,ti,bm in FIG_LIST: blocks_to_insert.append(list_entry(lbl,ti,bm))
for el in blocks_to_insert: abstract.addprevious(el)

print("== STEP 6: in-text reference fixes -> REF fields ==")
def replace_mention(find, bm, show):
    for p in body.iterchildren():
        if p.tag!=qn('w:p'): continue
        for r in p.findall(qn('w:r')):
            t=r.find(qn('w:t'))
            if t is not None and t.text and find in t.text:
                before,after=t.text.split(find,1); t.text=before
                anchor=r
                for f in field_runs(' REF %s \\h '%bm, show, sz=24): anchor.addnext(f); anchor=f
                anchor.addnext(make_run(after, sz=24)); return True
    return False
print(" 7.1->8.1:", replace_mention("Table 7.1","_Ref_tbl_8_1","Table 8.1"))
print(" 7.3->8.3:", replace_mention("Table 7.3","_Ref_tbl_8_3","Table 8.3"))
print(" 9.1 ref :", replace_mention("Table 9.1","_Ref_tbl_9_1","Table 9.1"))

print("== STEP 7: append references [16]-[18] ==")
def ref_p(text):
    p=mk('p'); ppr=mk('pPr')
    sp=mk('spacing'); setval(sp,'after','120'); setval(sp,'line','240'); setval(sp,'lineRule','auto'); ppr.append(sp)
    jc=mk('jc'); setval(jc,'val','both'); ppr.append(jc)
    p.append(ppr); p.append(make_run(text, sz=22)); return p
ref15=find_by_text("[15] Yao, S., Yu, D., Zhao")
new_refs=[
'[16] Jimenez, C. E., Yang, J., Wettig, A., Yao, S., Pei, K., Press, O., & Narasimhan, K., “SWE-bench: Can Language Models Resolve Real-World GitHub Issues?”, International Conference on Learning Representations (ICLR), 2024.',
'[17] Aleithan, R., Xue, H., Mohajer, M. M., Nnorom, E., Uddin, G., & Wang, S., “SWE-Bench+: Enhanced Coding Benchmark for LLMs”, arXiv:2410.06992, 2024.',
'[18] Huang, J., Chen, X., Mishra, S., Zheng, H. S., Yu, A. W., Song, X., & Zhou, D., “Large Language Models Cannot Self-Correct Reasoning Yet”, International Conference on Learning Representations (ICLR), 2024.']
anchor=ref15
for rt in new_refs:
    rp=ref_p(rt); anchor.addnext(rp); anchor=rp

print("== STEP 8: Resume / Ozgecmis section at end ==")
sectPr=body.find(qn('w:sectPr'))
def end_insert(el): sectPr.addprevious(el)
end_insert(title_heading("RESUME"))
def resume_block(name, sid, email):
    lines=[("Name Surname: ", name),("Student Number: ", sid),
           ("Department: ", "Department of Software Engineering, Faculty of Engineering, Fatih Sultan Mehmet Vakif University"),
           ("E-mail: ", email),
           ("Place / Date of Birth: ", "[ ... ]"),
           ("Bachelor's Degree: ", "Fatih Sultan Mehmet Vakif University, Department of Software Engineering, [graduation year]"),
           ("Work Experience / Awards / Publications: ", "[optional - to be completed]")]
    for k,v in lines:
        p=mk('p'); ppr=mk('pPr')
        sp=mk('spacing'); setval(sp,'after','60'); setval(sp,'line','240'); setval(sp,'lineRule','auto'); ppr.append(sp)
        jc=mk('jc'); setval(jc,'val','both'); ppr.append(jc); p.append(ppr)
        p.append(make_run(k, bold=True, sz=24)); p.append(make_run(v, sz=24))
        end_insert(p)
    end_insert(mk('p'))
resume_block("Ayşe Serra GÜMÜŞTAKIM","2121251008","[ayse.gumustakim]@stu.fsm.edu.tr")
resume_block("Ömer Yasir ÖNAL","2221221562","omeryasir.onal@stu.fsm.edu.tr")

# NOTE: The original CONTENTS is already a native Word Table-of-Contents field
# (wrapped in a w:sdt content control). It auto-updates in Word: once the figure
# and table captions are no longer heading-styled (done above), updating the TOC
# removes the stale "Figure/Table ..." sub-lines and corrects page numbers.
# We therefore leave it untouched rather than adding a second TOC.

doc.save(OUT)
print("SAVED ALL ->", OUT)
