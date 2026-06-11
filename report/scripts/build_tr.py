#!/usr/bin/env python3
"""Produce the fully-Turkish report from the English deliverable.

The English file and the Turkish file share an identical paragraph/table
structure (same indices, images, bookmarks, TOC and PAGEREF fields); this
script overwrites the text of each translatable unit with its Turkish
equivalent from tr_translation_map.py, preserving all formatting and fields.
Run from the repo root:  python report/scripts/build_tr.py
"""
import os, sys, shutil, zipfile
HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
from tr_translation_map import TR, LIST, REFPARAS, TABLES
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

SRC = os.path.join(HERE, "..", "AKIS_Graduation_Report_EN.docx")
OUT = os.path.join(HERE, "..", "AKIS_Bitirme_Raporu_TR.docx")
shutil.copy(SRC, OUT)
d = Document(OUT); P = d.paragraphs

def is_img(r): return bool(r._element.findall('.//'+qn('a:blip')))
def set_text(p, t):
    tr = [r for r in p.runs if not is_img(r)]
    if not tr:
        p.add_run(t); return
    tr[0].text = t
    for r in tr[1:]: r.text = ''

for i, v in TR.items(): set_text(P[i], v)
for i, v in LIST.items():
    if P[i].runs: P[i].runs[0].text = v
for i, v in REFPARAS.items():
    p = P[i]; ref = p.runs[0] if p.runs else None
    name = (ref.font.name if ref else None) or 'Times New Roman'
    size = ref.font.size if ref else None
    for ch in list(p._p):
        if ch.tag in (qn('w:r'), qn('w:fldSimple'), qn('w:hyperlink')): p._p.remove(ch)
    r = p.add_run(v); r.font.name = name; r.font.size = size or Pt(12)
for ti, rows in TABLES.items():
    t = d.tables[ti]
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            set_text(t.rows[ri].cells[ci].paragraphs[0], val)
            for ex in t.rows[ri].cells[ci].paragraphs[1:]:
                for r in ex.runs: r.text = ''
d.save(OUT)
assert zipfile.ZipFile(OUT).testzip() is None
print("wrote", os.path.normpath(OUT))
