#!/usr/bin/env python3
"""Produce the clean English deliverable: replace spaced em/en dashes and
spaced hyphens between words with report-grade punctuation (parentheses for
appositive pairs, colons for label lists, commas for contrasts). The
non-spaced dash in the title is intentionally left untouched.

SRC is the prior combined report (English body + Turkish ÖZET), preserved in
git history at commit d2b2b4d. Restore it there to re-run:
    git show d2b2b4d:report/AKIS_bitirme_raporu_final.docx > /tmp/base.docx
"""
import os, sys, re, shutil
HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
from declash import declash
from docx import Document
from docx.oxml.ns import qn

SRC = os.environ.get("BASE_DOCX", "/tmp/base.docx")
OUT = os.path.join(HERE, "..", "AKIS_Graduation_Report_EN.docx")
shutil.copy(SRC, OUT)
d = Document(OUT)

def has_field(p):
    return any(r._element.findall('.//'+qn('w:fldChar')) or r._element.findall('.//'+qn('w:instrText')) for r in p.runs)
def has_img(p):
    return any(r._element.findall('.//'+qn('a:blip')) for r in p.runs)

changes = []
def process_para(p, idx):
    full = p.text
    new_full = declash(full)
    if new_full == full:
        return
    # try per-run (preserves formatting) if dashes don't cross run boundaries
    rts = [r.text for r in p.runs]
    new_rts = [declash(rt) for rt in rts]
    if ''.join(new_rts) == new_full:
        for r, nt in zip(p.runs, new_rts):
            if r.text != nt:
                r.text = nt
    elif not has_field(p) and not has_img(p):
        # collapse to run[0] (uniform body prose)
        p.runs[0].text = new_full
        for r in p.runs[1:]:
            r.text = ''
    else:
        changes.append((idx, "SKIPPED-FIELD", full, new_full))
        return
    changes.append((idx, "OK", full[:60], new_full[:60]))

for i,p in enumerate(d.paragraphs):
    process_para(p, i)
# tables
for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                process_para(p, 'tbl')

# --- Cover author lines: remove spaced hyphen between name and student number ---
def fix_cover(p):
    changed=False
    for r in p.runs:
        if r.text and re.search(r'[A-Za-zÇĞİÖŞÜçğıöşü]\s+-\s+\d', r.text):
            r.text = re.sub(r'([A-Za-zÇĞİÖŞÜçğıöşü])\s+-\s+(\d)', r'\1   \2', r.text)
            changed=True
    return changed
cover_fixed=[]
for i,p in enumerate(d.paragraphs):
    if 'GÜMÜŞTAKIM' in p.text or 'ÖNAL' in p.text:
        if fix_cover(p): cover_fixed.append((i,p.text.replace(chr(10),' / ')))

d.save(OUT)
print(f"=== EN built: {OUT} | {len(changes)} dash-changes ===")
for idx,st,a,b in changes:
    print(f"  [{idx}] {st}")
    print(f"      - {a!r}")
    print(f"      + {b!r}")
print("\n=== cover author lines fixed ===")
for i,t in cover_fixed: print(f"  [{i}] {t!r}")
# integrity
import zipfile
print("\nzip OK:", zipfile.ZipFile(OUT).testzip() is None, "| paras:", len(Document(OUT).paragraphs))
# assert no spaced em/en dash remain in body prose (title excepted)
full=" ".join(p.text for p in Document(OUT).paragraphs)
print("remaining ' — ' :", full.count(' — '), "| ' – ':", full.count(' – '))
